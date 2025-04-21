import streamlit as st
import requests
from bs4 import BeautifulSoup
from transformers import pipeline
from fpdf import FPDF
import re
import io
from deep_translator import GoogleTranslator
from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime
import qrcode

# Sabit başlık (marka başlığı)
PROJECT_TITLE = "AI Web Sayfası Özeti"

# Dil seçimi eşlemesi
lang_map = {
    "Türkçe": None,
    "English": "en",
    "Deutsch (Almanca)": "de",
    "Français": "fr",
    "Español": "es",
    "Italiano": "it"
}

# PDF Sınıfı
class PDF(FPDF):
    def __init__(self):
        super().__init__()
        self.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
        self.add_font('DejaVu', 'B', 'DejaVuSans-Bold.ttf', uni=True)

    def header(self):
        self.image("synthesia_logo_light.png", x=10, y=8, w=30)
        self.set_xy(10, 25)
        self.set_font("DejaVu", style="B", size=14)
        self.cell(0, 10, PROJECT_TITLE, ln=True, align='L')
        self.ln(15)
        self.line(10, 35, 200, 35)

    def footer(self):
        self.set_y(-15)
        self.set_font("DejaVu", size=8)
        self.cell(0, 10, f"Sayfa {self.page_no()}", align='C')

st.set_page_config(page_title="AI Web Özetleyici", page_icon="🧠")





url = st.text_input("📎 Özetlemek istediğiniz web sayfasının URL'sini girin:", key="main_url")

with st.sidebar:
    st.image("synthesia_logo_dark.png", width=140)
    st.markdown("### Synthesia™")
    st.markdown("<p style='color: #cccccc;'>Yapay zekâ destekli özetleyiciye hoş geldiniz!</p>", unsafe_allow_html=True)

if not url:
    st.stop()
else:
    st.title(PROJECT_TITLE)
output_format = st.radio("📄 Çıktı formatı seçin:", ["PDF", "TXT", "DOCX"])
lang_choice = st.selectbox("🌍 Hangi dile çevrilsin? (isteğe bağlı)", list(lang_map.keys()))
selected_lang = lang_map[lang_choice]

if st.button("📝 Sayfayı Özetle"):
    if url:
        try:
            with st.spinner("Sayfa çekiliyor ve özetleniyor..."):
                response = requests.get(url)
                response.raise_for_status()
                soup = BeautifulSoup(response.text, 'html.parser')

                paragraphs = soup.find_all('p')
                metin = "".join(p.get_text() for p in paragraphs)

                if not metin.strip():
                    st.warning("Sayfada yeterli yazı içeriği bulunamadı.")
                else:
                    title_tag = soup.find('title')
                    meta_desc = soup.find('meta', attrs={'name': 'description'})
                    title_text = title_tag.get_text().strip() if title_tag else "ozet"
                    if meta_desc and meta_desc.get('content'):
                        title_text += " - " + meta_desc['content'].strip()
                    file_base = re.sub(r'\W+', '_', title_text.lower())[:30]

                    summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
                    chunks = [metin[i:i+1000] for i in range(0, len(metin), 1000)]
                    summaries = [summarizer(chunk, do_sample=False)[0]['summary_text'] for chunk in chunks]
                    final_summary = "\n".join(summaries)

                    if selected_lang:
                        try:
                            title_text = GoogleTranslator(source='auto', target=selected_lang).translate(title_text)
                            final_summary = GoogleTranslator(source='auto', target=selected_lang).translate(final_summary)
                            source_label = GoogleTranslator(source='auto', target=selected_lang).translate("Bu özetin kaynağı:")
                        except Exception as e:
                            st.error(f"❌ Çeviri yapılamadı: {str(e)}")
                            source_label = "Bu özetin kaynağı:"
                    else:
                        source_label = "Bu özetin kaynağı:"

                    st.subheader("📌 Özet:")
                    st.write(final_summary)

                    filename = f"{file_base}_ozet.{output_format.lower()}"

                    if output_format == "PDF":
                        pdf = PDF()
                        pdf.add_page()
                        pdf.set_font("DejaVu", size=12)
                        pdf.multi_cell(0, 10, f"{title_text}\n\n{final_summary}")
                        qr = qrcode.make(url)
                        qr_path = "qr_temp.png"
                        qr.save(qr_path)
                        pdf.image(qr_path, x=160, y=pdf.get_y()+10, w=30)
                        pdf.ln(35)
                        pdf.cell(0, 10, f"{source_label} {url}", ln=True)
                        pdf_output = pdf.output(dest='S')
                        pdf_buffer = io.BytesIO(pdf_output)
                        st.download_button("📥 Özeti indir (PDF)", pdf_buffer, file_name=filename)

                    elif output_format == "TXT":
                        buffer = io.StringIO()
                        buffer.write(f"{title_text}\n\n{final_summary}")
                        st.download_button("📥 Özeti indir (TXT)", buffer.getvalue(), file_name=filename)

                    elif output_format == "DOCX":
                        doc = Document()

                        # Kapak hissiyatı veren giriş kısmı
                        doc.add_picture("synthesia_logo_light.png", width=Inches(2.5))
                        doc.add_heading(PROJECT_TITLE, level=1)
                        doc.add_paragraph(f"Kaynak: {title_text}")
                        doc.add_paragraph(f"Dil: {lang_choice}   Tarih: {datetime.today().strftime('%d.%m.%Y')}")
                        doc.add_paragraph("\n" + "-" * 40 + "\n")

                        doc.add_paragraph(final_summary)
                        doc.add_paragraph("\n" + source_label + " " + url)

                        docx_buffer = io.BytesIO()
                        doc.save(docx_buffer)
                        docx_buffer.seek(0)
                        st.download_button("📥 Özeti indir (DOCX)", docx_buffer, file_name=filename)

                    st.success(f"Özet başarıyla oluşturuldu ve indirilmeye hazır: `{filename}`")

        except Exception as e:
            st.error(f"Hata oluştu: {str(e)}")
    else:
        st.warning("Lütfen bir URL girin.")
